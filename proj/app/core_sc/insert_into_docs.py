import docx #установить еще python-docx
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from datetime import date, datetime
from django.conf import settings

def create_docx(number_act, town, name_origin, number_contract, date_contract, project_original, name_of_work, result_of_work, date_start, date_end, stavka, cost, time_of_work, type_of_act):
    
    #doc = docx.Document("/template_s.docx")
    # Путь к файлу относительно BASE_DIR
    file_path = os.path.join(settings.BASE_DIR, 'app/core_sc/template_s.docx')
    
    # Проверка существования файла
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден по пути: {file_path}")
    
    #file_path = os.path.join(os.path.dirname(__file__), 'templates', 'template_s.docx')
    doc = docx.Document(file_path)
    #mas_tables=[]
    #перебор по ячейкам таблиц
    """ for table in doc.tables:
        mas_tables.append(table)
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text) """

    #print(len(mas_tables))
    table = doc.tables[0]
    row = table.rows[0]
    cell = row.cells[0]

    cell.text = f'Акт выполненных работ № {number_act}'
    #выравнивание по центру по горизонтали
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #----------------------------Город и дата
    table = doc.tables[1]
    row = table.rows[0]
    cell = row.cells[0]
    cell.text = f'г. {town}'

    cell = row.cells[1]
    today_date = date.today() 
    currentDate = datetime.strftime(today_date, "%d-%m-%y")  # сегодняшняя дата
    cell.text = f'{currentDate} г.'
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    #-----------------------------Стороны и договор
    table = doc.tables[2]
    row = table.rows[0]
    cell = row.cells[0]
    #cell.text = f'ООО «Интеллектуальные решения», именуемое в дальнейшем «Заказчик», в лице Генерального директора Киракосяна Левона Хачатуровича, действующего на основании Устава, с одной стороны, и ИП {name_origin}, именуемый в дальнейшем «Исполнитель», с другой стороны, именуемые совместно «Стороны», составили настоящий акт о нижеследующем: n/n/ В рамках Договора № {number_contract} на оказание услуг от {date_contract} г., Исполнителем были выполнены следующие работы по Проекту:'
    if type_of_act == 0:
        par1 = cell.add_paragraph().add_run(f'ООО «Интеллектуальные решения», именуемое в дальнейшем «Заказчик», в лице Генерального директора Киракосяна Левона Хачатуровича, действующего на основании Устава, с одной стороны, и ИП {name_origin}, именуемый в дальнейшем «Исполнитель», с другой стороны, именуемые совместно «Стороны», составили настоящий акт о нижеследующем:')
    else:
        par1 = cell.add_paragraph().add_run(f'ООО «Интеллектуальные решения», именуемое в дальнейшем «Заказчик», в лице Генерального директора Киракосяна Левона Хачатуровича, действующего на основании Устава, с одной стороны, и {name_origin}, именуемый в дальнейшем «Исполнитель», с другой стороны, именуемые совместно «Стороны», составили настоящий акт о нижеследующем:')
    par2 = cell.add_paragraph('')
    par2 = cell.add_paragraph().add_run(f'В рамках Договора № {number_contract} на оказание услуг от {date_contract} г., Исполнителем были выполнены следующие работы по Проекту:')

    #-----------------------------таблица
    table = doc.tables[3]
    row = table.rows[1]
    cell = row.cells[1]
    par1 = cell.add_paragraph().add_run(f'{name_of_work}')
    par1.font.size = Pt(9) #размер шрифта
    par2 = cell.add_paragraph('')
    par2 = cell.add_paragraph().add_run(f'Проект: {project_original}')
    par2.font.size = Pt(9)

    cell = row.cells[2]
    par1 = cell.add_paragraph().add_run(f'{result_of_work}')
    par1.font.size = Pt(9) #размер шрифта

    cell = row.cells[3]
    par1 = cell.add_paragraph().add_run(f'С {date_start} по {date_end}')
    par1.font.size = Pt(9) #размер шрифта

    cell = row.cells[4]
    par1 = cell.add_paragraph().add_run(f'{time_of_work}')
    par1.font.size = Pt(9) #размер шрифта

    cell = row.cells[5]
    par1 = cell.add_paragraph().add_run(f'{stavka}')
    par1.font.size = Pt(9) #размер шрифта

    cell = row.cells[6]
    par1 = cell.add_paragraph().add_run(f'{cost}')
    par1.font.size = Pt(9) #размер шрифта

    row = table.rows[2]
    cell = row.cells[6]
    par1 = cell.add_paragraph().add_run(f'Итого: {cost}')
    par1.font.size = Pt(9) #размер шрифта

    #-----------------------------текс после таблицы
    table = doc.tables[4]
    row = table.rows[0]
    cell = row.cells[0]
    cost_separate = cost.split('.')
    cost_part1 = cost_separate[0]
    cost_part2 = cost_separate[1]
    cell.text = f'Работы выполнены надлежащим образом, своевременно, качество выполненных работ соответствует условиям Договора. Претензий со стороны Заказчика к Исполнителю не имеется. Стоимость выполненных работ составляет сумму {cost_part1} руб. {cost_part2} коп.'

    #-----------------------------подписи на акте
    table = doc.tables[5]
    row = table.rows[0]
    cell = row.cells[0]
    par1 = cell.add_paragraph().add_run('Заказчик')
    par2 = cell.add_paragraph().add_run('ООО "Интеллектуальные решения"')
    cell = row.cells[1]
    par1 = cell.add_paragraph().add_run('Исполнитель')
    if type_of_act == 0:
        par2 = cell.add_paragraph().add_run(f'ИП {name_origin}')
    else:
        par2 = cell.add_paragraph().add_run(f'{name_origin}')
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    #-----------------------------номер акта в завке
    table = doc.tables[7]
    row = table.rows[0]
    cell = row.cells[0]
    cell.text = f'Заявка на оказание услуг №{number_act}'
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #-----------------------------номер договора в завке
    table = doc.tables[8]
    row = table.rows[0]
    cell = row.cells[0]
    cell.text = f'по договору № {number_contract} от {date_contract}'
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #-----------------------------таблица в заявке
    table = doc.tables[9]
    row = table.rows[1]
    cell = row.cells[0]
    cell.text = f'{name_of_work}'
    cell = row.cells[1]
    cell.text = f'{stavka}'
    cell = row.cells[2]
    cell.text = f'{date_end}'
    cell = row.cells[3]
    cell.text = f'{cost}'

    #-----------------------------подписи в заявке
    table = doc.tables[10]
    row = table.rows[0]
    cell = row.cells[0]
    par1 = cell.add_paragraph().add_run('Заказчик')
    par2 = cell.add_paragraph().add_run('ООО "Интеллектуальные решения"')
    cell = row.cells[1]
    par1 = cell.add_paragraph().add_run('Исполнитель')
    if type_of_act == 0:
        par2 = cell.add_paragraph().add_run(f'ИП {name_origin}')
    else:
        par2 = cell.add_paragraph().add_run(f'{name_origin}')
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    path_act = os.path.join(settings.BASE_DIR, f'app/core_sc/documents/{name_origin}-Акт №{number_act} от {currentDate}.docx')
    #path_act = os.path.join(settings.BASE_DIR, '/app/templates/documents{name_origin}-Акт №{number_act}.docx')
    #path_act = "/proj/app/core_sc/template1.docx"
    #doc.save("d:\\act_bot_v3\\acts\\testdocx.docx")
    
    doc.save(path_act)
    return path_act